from __future__ import annotations
# --- path bootstrap for imports ---
import sys
from pathlib import Path
BASE = Path(__file__).resolve().parents[1]
if str(BASE) not in sys.path:
    sys.path.insert(0, str(BASE))
# -----------------------------------

import streamlit as st
import pandas as pd
import io, csv

from src.app.api import run_generation
from src.app.utils import load_yaml
from ui.utils import (
    _read_csv_robust, normalize_calendar,
    load_normalization_rules, save_normalization_rules,
    load_calendar, list_personas, df_to_markdown_row,
    discover_with_reasons, upsert_registry, get_registry, display_name_for
)
from ui.components.post_preview import render_post
from ui.components.sidebar import sidebar_controls

st.set_page_config(page_title="LinkedIn Post Generator (Minimal UI)", layout="wide")
st.title("LinkedIn Post Generator (Minimal UI)")

st.caption(f"Project root: {BASE}")
st.caption(f"Data folder: {BASE / 'data'}")

personas_yaml = load_yaml(str(BASE / "config" / "personas.yaml"))
personas = list_personas(personas_yaml)

valid, invalid = discover_with_reasons()
registry = get_registry()

col_tools = st.columns([1,1,3])[0]
with col_tools:
    if st.button("Refresh calendars"):
        valid, invalid = discover_with_reasons()

uploaded = st.file_uploader("Upload a schedule CSV (optional)", type=["csv"])
if uploaded:
    temp_path = BASE / "data" / uploaded.name
    with open(temp_path, "wb") as f:
        f.write(uploaded.getbuffer())
    valid, invalid = discover_with_reasons()

normalize_flag = st.checkbox("Normalize calendars on load (recommended)", value=True)

# ---- Manual Mapper ----
with st.sidebar.expander("Debug: discovered calendars", expanded=True):
    st.write("•")
    for fn in valid:
        st.write(fn)
    st.write("---")
    st.write("Ignored (not calendar-like or missing topic):")
    for fn in invalid.keys():
        st.write("•", fn)

with st.sidebar.expander("Fix an ignored calendar"):
    if invalid:
        pick = st.selectbox("Choose file", list(invalid.keys()))
        if pick:
            raw_df = _read_csv_robust(BASE / "data" / pick, nrows=5)
            st.caption(f"Columns detected: {list(raw_df.columns)}")
            cols = ["(none)"] + list(raw_df.columns)
            rules = load_normalization_rules()
            files_map = rules.setdefault("files", {}).setdefault(pick, {})
            aliases = files_map.setdefault("column_aliases", {})
            # Five canonical fields
            date_sel = st.selectbox("Map to 'date'", cols, index=cols.index(next((c for c in cols if c.lower() in ['date','day','publish date','post date','(none)']), '(none)')))
            topic_sel = st.selectbox("Map to 'topic'", cols, index=cols.index(next((c for c in cols if c.lower() in ['topic','hook theme','title','theme','(none)']), '(none)')))
            service_sel = st.selectbox("Map to 'service'", cols, index=cols.index(next((c for c in cols if 'service' in c.lower() or c.lower() in ['ardent service tie-in','(none)']), '(none)')))
            pillar_sel = st.selectbox("Map to 'pillar'", cols, index=cols.index(next((c for c in cols if 'pillar' in c.lower() or c.lower()=='content pillar' or c.lower()=='(none)']), '(none)')))
            audience_sel = st.selectbox("Map to 'audience'", cols, index=cols.index(next((c for c in cols if 'audience' in c.lower() or 'industry' in c.lower() or c.lower()=='(none)']), '(none)')))
            if st.button("Save mapping & refresh"):
                for src, tgt in [(date_sel,'date'), (topic_sel,'topic'), (service_sel,'service'), (pillar_sel,'pillar'), (audience_sel,'audience')]:
                    if src and src != "(none)":
                        aliases[src] = tgt
                save_normalization_rules(rules)
                st.success("Saved. Click 'Refresh calendars' at the top.")
    else:
        st.caption("All good — nothing to fix.")

if valid:
    from ui.components.sidebar import sidebar_controls
    controls = sidebar_controls(personas, valid, registry, invalid)
    sel_file = controls["schedule"]
    with st.sidebar.expander("Calendar name (alias)"):
        current_alias = registry.get(sel_file, sel_file)
        new_alias = st.text_input("Display name", value=current_alias, help="Alias for this calendar")
        if st.button("Save name"):
            upsert_registry(sel_file, new_alias)
            st.success("Saved")
else:
    st.info("No valid calendars found in /data. Map the columns for any ignored file and refresh.")
    controls = {"persona":"ardent_v2","schedule":"", "regenerate":False, "run": False}
    sel_file = ""

with st.expander("Calendar diagnostics", expanded=not bool(valid)):
    st.write("**Valid calendars:**", valid or "None")
    st.write("**Invalid calendars (with reasons):**")
    if invalid:
        for name, reason in invalid.items():
            st.write(f"- **{name}** → {reason}")
    else:
        st.write("None detected.")

col_left, col_right = st.columns([1,2])

with col_left:
    st.subheader("Schedule")
    if sel_file:
        try:
            from ui.utils import _read_csv_robust as read_csv
            raw_df = read_csv((BASE / "data" / sel_file))
            df = normalize_calendar(raw_df, sel_file) if normalize_flag else raw_df

            if "date" in df.columns:
                dt = pd.to_datetime(df["date"], errors="coerce")
                df = df.assign(date=dt)
                df = df.sort_values("date", na_position="last")
            st.caption(f"Using: **{display_name_for(sel_file)}** ({sel_file})")
            st.dataframe(df, width='stretch', hide_index=True)

            if st.button("Download normalized CSV"):
                buf = io.StringIO()
                out_df = df.copy()
                if "date" in out_df.columns and pd.api.types.is_datetime64_any_dtype(out_df["date"]):
                    out_df["date"] = out_df["date"].dt.strftime("%Y-%m-%d")
                out_df.to_csv(buf, index=False)
                st.download_button("Save file", data=buf.getvalue(), file_name=f"normalized_{sel_file}", mime="text/csv")

            if not df.empty:
                default_idx = int(df.index[0])
                if "date" in df.columns and pd.api.types.is_datetime64_any_dtype(df["date"]) and df["date"].notna().any():
                    future = df[df["date"] >= pd.Timestamp.today().normalize()]
                    if not future.empty:
                        default_idx = int(future.index[0])
                pick_idx = st.number_input("Row to generate",
                                           min_value=int(df.index.min()),
                                           max_value=int(df.index.max()),
                                           value=default_idx, step=1)
                sel = df.loc[pick_idx].to_dict()
                if isinstance(sel.get("date"), pd.Timestamp):
                    sel["date"] = sel["date"].strftime("%Y-%m-%d")
                st.markdown(df_to_markdown_row(sel))
            else:
                sel = {}
        except Exception as e:
            st.error(f"Failed to load calendar: {e}")
            sel = {}
    else:
        sel = {}

with col_right:
    st.subheader("Generation")
    if controls.get("run") and sel:
        topic = str(sel.get("topic","")).strip() or "Untitled Topic"
        service = [str(sel.get("service","")).strip()] if sel.get("service") else []
        payload = run_generation(topic=topic, services=service, persona_key=controls["persona"])
        render_post(payload)

        st.divider()
        st.subheader("Export")
        csv_rows = [{
            "date": sel.get("date",""),
            "persona": controls["persona"],
            "topic": topic,
            "service": ",".join(service),
            "body": payload.get("body",""),
            "hashtags": " ".join(payload.get("hashtags",[])),
            "sources": "\n".join(payload.get("sources",[])),
        }]
        csv_buf = io.StringIO()
        writer = csv.DictWriter(csv_buf, fieldnames=list(csv_rows[0].keys()))
        writer.writeheader(); writer.writerows(csv_rows)
        st.download_button("Download CSV", data=csv_buf.getvalue(),
                           file_name="linkedin_post.csv", mime="text/csv")

        try:
            from docx import Document
            doc = Document()
            doc.add_heading("LinkedIn Post", level=1)
            doc.add_paragraph(payload.get("body",""))
            doc.add_paragraph("")
            doc.add_paragraph("Hashtags: " + " ".join(payload.get("hashtags",[])))
            if payload.get("sources"):
                doc.add_paragraph("Sources:")
                for s in payload["sources"]:
                    doc.add_paragraph(s)
            import tempfile
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp.name)
            with open(tmp.name, "rb") as f:
                st.download_button("Download DOCX", data=f.read(),
                                   file_name="linkedin_post.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception:
            st.info("Install python-docx to enable DOCX exports.")
    else:
        st.write("Pick a row on the left and click **Generate Post** in the sidebar.")
